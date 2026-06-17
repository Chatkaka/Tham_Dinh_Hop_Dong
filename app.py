import io
import os
import re
import requests
import pandas as pd
import streamlit as st
from fpdf import FPDF
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from utils.docx_parser import doc_hop_dong_toan_dien
# from utils.appraiser import appraise_contract

# Cấu hình trang Streamlit
st.set_page_config(
    page_title="App Thẩm định Hợp đồng Mua bán",
    page_icon="📝",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Nhúng CSS tùy chỉnh để tối ưu hóa trải nghiệm giao diện người dùng
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
    
    /* Font mặc định cho toàn bộ trang */
    html, body, [class*="css"] {
        font-family: 'Inter', sans-serif;
    }
    
    /* Thiết kế Header cao cấp */
    .header-container {
        background: linear-gradient(135deg, #1e3a8a 0%, #0f766e 100%);
        padding: 2.2rem;
        border-radius: 12px;
        color: white;
        margin-bottom: 2rem;
        box-shadow: 0 4px 20px rgba(0, 0, 0, 0.08);
        text-align: center;
    }
    
    .header-title {
        font-size: 2.4rem;
        font-weight: 700;
        margin-bottom: 0.4rem;
        color: #ffffff;
    }
    
    .header-subtitle {
        font-size: 1.1rem;
        font-weight: 300;
        color: #e2e8f0;
    }
    
    /* Hộp chứa văn bản hợp đồng */
    .contract-viewer {
        background-color: #ffffff;
        border: 1px solid #e2e8f0;
        padding: 1.8rem;
        border-radius: 12px;
        height: 700px;
        overflow-y: auto;
        font-size: 14.5px;
        line-height: 1.7;
        color: #334155;
        box-shadow: inset 0 2px 8px rgba(0,0,0,0.02);
    }
    
    /* Định dạng thanh cuộn cho hợp đồng */
    .contract-viewer::-webkit-scrollbar {
        width: 6px;
    }
    .contract-viewer::-webkit-scrollbar-track {
        background: #f1f5f9;
    }
    .contract-viewer::-webkit-scrollbar-thumb {
        background: #cbd5e1;
        border-radius: 4px;
    }
    .contract-viewer::-webkit-scrollbar-thumb:hover {
        background: #94a3b8;
    }
    
    /* Thẻ Card cho Dashboard tóm tắt */
    .dashboard-card {
        background-color: #ffffff;
        border: 1px solid #e2e8f0;
        border-radius: 10px;
        padding: 1.2rem;
        box-shadow: 0 2px 5px rgba(0,0,0,0.01);
        text-align: center;
    }
    .card-value {
        font-size: 1.8rem;
        font-weight: 700;
        margin-top: 0.3rem;
    }
    .card-label {
        font-size: 0.9rem;
        color: #64748b;
        font-weight: 500;
        text-transform: uppercase;
        letter-spacing: 0.5px;
    }
    
    /* Viền chỉ định mức rủi ro */
    .risk-high { border-top: 5px solid #ef4444; }
    .risk-medium { border-top: 5px solid #f59e0b; }
    .risk-low { border-top: 5px solid #10b981; }
    .risk-total { border-top: 5px solid #3b82f6; }
    
    /* Vùng kết luận tổng hợp */
    .summary-verdict {
        padding: 1.2rem;
        border-radius: 8px;
        margin-bottom: 1.5rem;
        font-weight: 500;
        display: flex;
        align-items: center;
        gap: 10px;
        box-shadow: 0 2px 4px rgba(0,0,0,0.01);
    }
    
    .verdict-danger {
        background-color: #fef2f2;
        color: #991b1b;
        border-left: 5px solid #ef4444;
    }
    
    .verdict-warning {
        background-color: #fffbeb;
        color: #92400e;
        border-left: 5px solid #f59e0b;
    }
    
    .verdict-success {
        background-color: #ecfdf5;
        color: #065f46;
        border-left: 5px solid #10b981;
    }
    
    /* Style cho các nút Export tải về */
    .stDownloadButton > button {
        background: linear-gradient(135deg, #0f766e 0%, #0d9488 100%) !important;
        color: white !important;
        border: none !important;
        font-weight: 600 !important;
        transition: transform 0.2s, box-shadow 0.2s !important;
    }
    .stDownloadButton > button:hover {
        transform: scale(1.02) !important;
        box-shadow: 0 4px 10px rgba(13, 148, 136, 0.3) !important;
    }
    
</style>
""", unsafe_allow_html=True)



# Helper: Đọc các file hợp đồng mẫu để làm tiêu chuẩn đối chiếu
@st.cache_data
def load_templates(selected_files=None):
    import os
    from docx import Document
    template_dir = "templates"
    os.makedirs(template_dir, exist_ok=True)
    template_texts = []
    
    if not selected_files:
        return "Không có hợp đồng mẫu nào được chọn để đối chiếu."
        
    for filename in selected_files:
        filepath = os.path.join(template_dir, filename)
        if not os.path.exists(filepath):
            continue
            
        if filename.endswith(".docx"):
            try:
                doc = Document(filepath)
                text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                template_texts.append(f"--- Mẫu: {filename} ---\n{text}")
            except Exception as e:
                print(f"Lỗi đọc {filename}: {e}")
        elif filename.endswith(".txt"):
            try:
                with open(filepath, 'r', encoding='utf-8') as f:
                    template_texts.append(f"--- Mẫu: {filename} ---\n{f.read()}")
            except Exception as e:
                print(f"Lỗi đọc {filename}: {e}")
                
    if not template_texts:
        return "Chưa có file hợp đồng mẫu hợp lệ nào."
        
    return "\n\n".join(template_texts)

# Helper: Gọi Gemini API để thẩm định hợp đồng
def tham_dinh_bang_gemini(noi_dung_hop_dong, api_key, tieu_chuan_doi_chieu):
    try:
        from google import genai
    except ImportError:
        return "Lỗi: Thư viện google-genai chưa được cài đặt. Vui lòng cài đặt trước."
        
    try:
        client = genai.Client(api_key=api_key)
        prompt_chuyen_gia = f"""
Mục tiêu:
Bạn là một Chuyên gia Pháp lý và Luật sư doanh nghiệp dày dạn kinh nghiệm. Nhiệm vụ của bạn là thẩm định chi tiết các điều khoản trong văn bản hợp đồng được cung cấp, dựa trên cơ sở pháp luật hiện hành của Việt Nam.

Bối cảnh:
Văn bản đầu vào là một hợp đồng pháp lý. Các điều khoản trong hợp đồng cần phải rõ ràng, thống nhất, không được có sự xung đột nội bộ giữa các điều khoản với nhau, và tuyệt đối không được trái với các quy định của pháp luật Việt Nam.

Hướng dẫn thực hiện (Chuỗi suy luận từng bước):
Hãy thực hiện nhiệm vụ này một cách hệ thống theo các bước sau. Đối với mỗi bước, hãy giải thích quá trình suy luận của bạn:

1. Phân tích tổng quan: Đọc toàn bộ nội dung hợp đồng để hiểu rõ ngữ cảnh, đối tượng và mục đích của hợp đồng.
2. Kiểm tra tính thống nhất nội bộ: Đối chiếu chéo các điều khoản, câu từ và định nghĩa trong hợp đồng. Xác định và liệt kê ra bất kỳ sự mâu thuẫn, xung đột hoặc mơ hồ nào giữa các điều khoản.
3. Thẩm định tính tuân thủ pháp luật: Đánh giá từng điều khoản chính dựa trên các nguyên tắc cơ bản của pháp luật Việt Nam (Dân sự, Thương mại, Lao động... tùy thuộc vào bản chất hợp đồng). Xác định bất kỳ điều khoản nào có nguy cơ vi phạm pháp luật hoặc vô hiệu.
4. Đối chiếu Hợp đồng mẫu (nếu có): So sánh các điều khoản trong hợp đồng tải lên với 'Tiêu chuẩn đối chiếu' (Hợp đồng mẫu) dưới đây. Chỉ ra bất kỳ sự sai lệch, thiếu sót hoặc thay đổi nào gây bất lợi.
5. Đề xuất chỉnh sửa: Đối với mỗi vấn đề được phát hiện ở Bước 2, Bước 3 và Bước 4, hãy cung cấp các khuyến nghị cụ thể hoặc đoạn văn bản viết lại để khắc phục sự cố, đảm bảo tính chặt chẽ và hợp pháp.

Giới hạn (Constraints):
- Trình bày phản hồi của bạn dưới dạng danh sách gạch đầu dòng (bulleted list) rõ ràng. KHÔNG trả về định dạng bảng.
- Sử dụng ngôn ngữ chuyên ngành pháp lý chính xác, khách quan và chuyên nghiệp.
- Chỉ tập trung vào văn bản được cung cấp. Nếu cần thêm thông tin hoặc bối cảnh từ người dùng để đánh giá một điều khoản, hãy nêu rõ câu hỏi bổ sung.
- Không được tự bịa đặt các điều luật không có thật.
- Nếu nội dung hợp đồng dưới đây trống hoặc không hợp lệ, trả lời duy nhất một câu: 'Lỗi: Không nhận được nội dung hợp đồng hoặc file tải lên không hợp lệ. Vui lòng kiểm tra lại.'

Tiêu chuẩn đối chiếu (Các file hợp đồng mẫu):
{tieu_chuan_doi_chieu}

Nội dung hợp đồng:
{noi_dung_hop_dong}
"""
        response = client.models.generate_content(
            model='gemini-3.5-flash',
            contents=prompt_chuyen_gia,
        )
        return response.text
    except Exception as e:
        return f"Lỗi gọi Gemini API: {str(e)}"

# Helper: Gọi Gemini API để chỉnh sửa hợp đồng theo yêu cầu
def chinh_sua_hop_dong_bang_gemini(noi_dung_hop_dong, api_key, yeu_cau_chinh_sua):
    try:
        from google import genai
    except ImportError:
        return "Lỗi: Thư viện google-genai chưa được cài đặt. Vui lòng cài đặt trước."
        
    try:
        client = genai.Client(api_key=api_key)
        prompt_chinh_sua = f"""
Mục tiêu:
Bạn là một Chuyên gia Pháp lý và Luật sư doanh nghiệp dày dạn kinh nghiệm. Nhiệm vụ của bạn là thực hiện chỉnh sửa, tối ưu hóa, viết lại hoặc bổ sung các điều khoản trong hợp đồng dựa trên yêu cầu cụ thể của người dùng, đảm bảo tính chặt chẽ về mặt pháp lý, hạn chế tối đa rủi ro cho các bên và sử dụng ngôn ngữ pháp lý chuyên nghiệp.

Nội dung hợp đồng gốc:
\"\"\"
{noi_dung_hop_dong}
\"\"\"

Yêu cầu chỉnh sửa của người dùng:
\"\"\"
{yeu_cau_chinh_sua}
\"\"\"

Hướng dẫn thực hiện:
1. Xác định điều khoản hoặc đoạn văn bản trong hợp đồng cần được chỉnh sửa dựa trên yêu cầu của người dùng.
2. Viết lại hoặc điều chỉnh điều khoản đó một cách chuyên nghiệp, chặt chẽ, tối ưu quyền lợi pháp lý và tuân thủ pháp luật Việt Nam hiện hành.
3. Giải thích ngắn gọn lý do thực hiện các chỉnh sửa đó (ví dụ: hạn chế rủi ro gì, tăng tính rõ ràng ở điểm nào).
4. Phản hồi của bạn cần được trình bày rõ ràng bằng tiếng Việt, dưới dạng Markdown, theo cấu trúc sau:
   - **Yêu cầu thực hiện**: Tóm tắt lại yêu cầu của người dùng.
   - **Nội dung điều khoản gốc** (nếu có): Trích dẫn lại đoạn văn bản gốc cần sửa.
   - **Nội dung đề xuất chỉnh sửa mới**: Đoạn văn bản hoàn chỉnh sau khi đã được viết lại/tối ưu.
   - **Giải thích & Phân tích pháp lý**: Giải thích ngắn gọn các điểm cải tiến và lợi ích pháp lý của phương án mới.
"""
        response = client.models.generate_content(
            model='gemini-3.5-flash',
            contents=prompt_chinh_sua,
        )
        return response.text
    except Exception as e:
        return f"Lỗi gọi Gemini API khi chỉnh sửa: {str(e)}"

# Helper: Phân tích bảng Markdown từ văn bản Gemini
def parse_markdown_table(md_text):
    lines = md_text.strip().split("\n")
    rows = []
    for line in lines:
        line = line.strip()
        if line.startswith("|") and line.endswith("|"):
            # Tách cột
            cells = [c.strip() for c in line.split("|")[1:-1]]
            # Bỏ qua dòng ngăn cách |---|---|
            is_separator = all(re.match(r'^[\s\-\:]+$', cell) for cell in cells) if cells else False
            if is_separator:
                continue
            rows.append(cells)
    return rows

# Helper: Xuất báo cáo AI sang tệp Word
def export_gemini_table_to_word(gemini_md, contract_name="Hợp đồng"):
    doc = Document()
    
    # Tiêu đề báo cáo
    title = doc.add_paragraph()
    run = title.add_run("BÁO CÁO THẨM ĐỊNH HỢP ĐỒNG BẰNG TRÍ TUỆ NHÂN TẠO (GEMINI)")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Arial'
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Thông tin chung
    sub = doc.add_paragraph()
    run_sub = sub.add_run(f"Tên văn bản: {contract_name}\nNgày thực hiện: {datetime.now().strftime('%d/%m/%Y')}\nMô hình AI: Gemini 3.5 Flash")
    run_sub.font.size = Pt(10)
    run_sub.font.name = 'Arial'
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph().add_run("\n")
    
    # Phân tích các dòng bảng
    rows = parse_markdown_table(gemini_md)
    if rows:
        headers = rows[0]
        data_rows = rows[1:]
        
        # Tạo bảng trong Word
        table = doc.add_table(rows=len(rows), cols=len(headers))
        table.style = 'Table Grid'
        
        # Định dạng Header
        hdr_cells = table.rows[0].cells
        for col_idx, text in enumerate(headers):
            hdr_cells[col_idx].text = text
            for p in hdr_cells[col_idx].paragraphs:
                for r in p.runs:
                    r.font.bold = True
                    r.font.name = 'Arial'
                    r.font.size = Pt(9.5)
                    
        # Ghi dữ liệu từng dòng
        for row_idx, cells in enumerate(data_rows):
            row_cells = table.rows[row_idx + 1].cells
            for col_idx, text in enumerate(cells):
                if col_idx < len(row_cells):
                    row_cells[col_idx].text = text
                    for p in row_cells[col_idx].paragraphs:
                        for r in p.runs:
                            r.font.name = 'Arial'
                            r.font.size = Pt(9.5)
    else:
        # Fallback ghi văn bản thuần nếu không tìm thấy bảng hợp lệ
        for line in gemini_md.split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip())
    doc_io = io.BytesIO()
    doc.save(doc_io)
    return doc_io.getvalue()




# Helper: Xuất báo cáo AI sang tệp PDF
def export_gemini_table_to_pdf(gemini_md, contract_name="Hợp đồng"):
    try:
        from fpdf import FPDF
    except ImportError:
        return b"FPDF2 not installed"
        
    class PDF(FPDF):
        def header(self):
            pass
            
    pdf = PDF(orientation='L', unit='mm', format='A4')
    pdf.add_page()
    
    try:
        pdf.add_font('Arial', '', r'C:\Windows\Fonts\arial.ttf')
        pdf.add_font('Arial', 'B', r'C:\Windows\Fonts\arialbd.ttf')
        font_name = 'Arial'
    except Exception:
        font_name = 'helvetica'
        
    pdf.set_font(font_name, 'B', 14)
    pdf.cell(0, 10, "BÁO CÁO THẨM ĐỊNH HỢP ĐỒNG BẰNG TRÍ TUỆ NHÂN TẠO (GEMINI)", new_x="LMARGIN", new_y="NEXT", align='C')
    pdf.set_font(font_name, '', 10)
    pdf.cell(0, 10, f"Tên văn bản: {contract_name} | Ngày thực hiện: {datetime.now().strftime('%d/%m/%Y')}", new_x="LMARGIN", new_y="NEXT", align='C')
    pdf.ln(5)
    
    rows = parse_markdown_table(gemini_md)
    if rows:
        with pdf.table(text_align="LEFT", width=277) as table:
            headers = rows[0]
            header_row = table.row()
            pdf.set_font(font_name, 'B', 9)
            for text in headers:
                header_row.cell(text)
                
            pdf.set_font(font_name, '', 9)
            for cells in rows[1:]:
                data_row = table.row()
                for text in cells:
                    data_row.cell(text)
    else:
        pdf.multi_cell(0, 10, gemini_md)
        
    return bytes(pdf.output())


# Giao diện chính ứng dụng
st.markdown("""
<div class="header-container">
    <div class="header-title">Trợ lý Thẩm định Hợp đồng Mua bán</div>
    <div class="header-subtitle">Hệ thống phân tích điều khoản, phát hiện rủi ro và tự động lập Báo cáo Thẩm định pháp lý</div>
</div>
""", unsafe_allow_html=True)

# Khởi tạo các trạng thái trong session_state
if "full_contract_text" not in st.session_state:
    st.session_state.full_contract_text = None

if "current_file_name" not in st.session_state:
    st.session_state.current_file_name = None
if "gemini_report" not in st.session_state:
    st.session_state.gemini_report = None
if "gemini_edit_result" not in st.session_state:
    st.session_state.gemini_edit_result = None

# Lấy khóa API mặc định từ secrets hoặc biến môi trường một cách an toàn
api_key = None
try:
    if "GEMINI_API_KEY" in st.secrets:
        api_key = st.secrets["GEMINI_API_KEY"]
except Exception:
    pass

if not api_key:
    api_key = os.environ.get("GEMINI_API_KEY")

uploaded_file = None

# Sidebar dùng để cấu hình API và tải file phụ trợ
with st.sidebar:
    st.image("https://img.icons8.com/clouds/100/law.png", width=80)
    st.subheader("Cấu hình & Tải file")
    uploaded_file_sidebar = st.file_uploader(
        "Tải lên Hợp đồng (.docx)",
        type=["docx"],
        key="sidebar_uploader",
        help="Chỉ hỗ trợ định dạng Word (.docx)"
    )
    if uploaded_file_sidebar:
        # Nếu file trong sidebar khác file đang hoạt động, load lại dữ liệu mới
        if st.session_state.current_file_name != uploaded_file_sidebar.name:
            st.session_state.current_file_name = uploaded_file_sidebar.name
            with st.spinner("Đang phân tích cấu trúc tài liệu..."):
                file_bytes = uploaded_file_sidebar.read()
                st.session_state.full_contract_text = doc_hop_dong_toan_dien(io.BytesIO(file_bytes))
                st.session_state.gemini_report = None
                st.session_state.gemini_edit_result = None
                st.toast("Tải tệp thành công!")
                st.rerun()
                
    st.divider()
    st.subheader("📁 Cập nhật Hợp đồng Mẫu")
    template_uploader = st.file_uploader(
        "Tải lên Hợp đồng Mẫu (.docx)",
        type=["docx"],
        key="template_uploader",
        help="Lưu đè và thay thế các file hợp đồng mẫu hiện tại dùng làm tiêu chuẩn thẩm định."
    )
    if template_uploader is not None:
        if st.session_state.get('last_uploaded_template') != template_uploader.name:
            import os
            template_dir = "templates"
            os.makedirs(template_dir, exist_ok=True)
            # Lưu file mới (không xóa file cũ)
            new_file_path = os.path.join(template_dir, template_uploader.name)
            with open(new_file_path, "wb") as f:
                f.write(template_uploader.getbuffer())
            st.session_state['last_uploaded_template'] = template_uploader.name
            st.success(f"Đã tải lên Hợp đồng mẫu: {template_uploader.name}")
            # Xóa cache để load_templates có thể được cập nhật (mặc dù giờ load_templates sẽ nhận đối số)
            load_templates.clear()

    # Tạo danh sách chọn Hợp đồng mẫu
    import os
    template_dir = "templates"
    available_templates = []
    if os.path.exists(template_dir):
        available_templates = [f for f in os.listdir(template_dir) if f.endswith(".docx") or f.endswith(".txt")]
        
    if available_templates:
        selected_templates = st.multiselect(
            "Chọn Hợp đồng mẫu để đối soát:",
            options=available_templates,
            default=available_templates,
            help="Hệ thống sẽ đối chiếu hợp đồng tải lên với các mẫu được chọn."
        )
        st.session_state.selected_templates = selected_templates
    else:
        st.session_state.selected_templates = []
            
    # Ô nhập API Key bổ sung nếu chưa có sẵn trên hệ thống
    if not api_key:
        st.divider()
        st.subheader("🔑 Gemini API Key")
        user_key = st.sidebar.text_input(
            "Nhập Gemini API Key:",
            type="password",
            help="Cần thiết để chạy tính năng thẩm định bằng mô hình AI Gemini 3.5 Flash."
        )
        if user_key:
            api_key = user_key
            


# ----------------- KHI CHƯA UPLOAD FILE -----------------
if not st.session_state.get("full_contract_text"):
    st.markdown("""
    <div style="background-color: #f8fafc; border: 1px dashed #cbd5e1; border-radius: 12px; padding: 2.5rem; text-align: center; margin-top: 1.5rem; margin-bottom: 1.5rem;">
        <img src="https://img.icons8.com/external-flatart-icons-outline-flatarticons/100/external-contract-business-and-finance-flatart-icons-outline-flatarticons.png" style="width: 80px; opacity: 0.8; margin-bottom: 1rem;"/>
        <h3 style="color: #1e3a8a; margin-bottom: 0.5rem;">Vui lòng tải lên tài liệu hợp đồng</h3>
        <p style="color: #64748b; max-width: 600px; margin: 0 auto 1rem auto;">
            Hệ thống hỗ trợ tải lên các tệp hợp đồng Word (.docx). Sau khi tải lên, hệ thống sẽ tự động quét văn bản gốc và lập bảng báo cáo đánh giá rủi ro pháp lý theo Luật Nhà ở và Luật Kinh doanh Bất động sản Việt Nam.
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    # Tạo giao diện upload file trực quan ngay trên màn hình chính
    uploaded_file_main = st.file_uploader(
        "Kéo thả hoặc chọn tệp hợp đồng ở đây (.docx):",
        type=["docx"],
        key="main_uploader",
        help="Chỉ hỗ trợ định dạng Word (.docx)"
    )
    if uploaded_file_main:
        st.session_state.current_file_name = uploaded_file_main.name
        with st.spinner("Đang phân tích cấu trúc tài liệu..."):
            file_bytes = uploaded_file_main.read()
            st.session_state.full_contract_text = doc_hop_dong_toan_dien(io.BytesIO(file_bytes))
            st.session_state.gemini_report = None
            st.session_state.gemini_edit_result = None
            st.toast("Tải tệp thành công!")
            st.rerun()
            
    st.markdown("---")
    
    # Hiển thị demo các tính năng chính dưới dạng các cột nhỏ
    st.write("### Các tính năng cốt lõi của ứng dụng:")
    col_f1, col_f2, col_f3 = st.columns(3)
    with col_f1:
        st.info("**👁️ Hiển thị trực quan**\n\nXem văn bản gốc ở cột trái được cấu trúc rõ ràng với các mục đề và bảng biểu.")
    with col_f2:
        st.warning("**⚡ Phát hiện rủi ro tự động**\n\nQuét điều khoản bảo hành, phạt vi phạm, sự kiện bất khả kháng và tiến độ thanh toán bằng thuật toán thông minh.")
    with col_f3:
        st.success("**✏️ Biên tập và Xuất báo cáo**\n\nChỉnh sửa báo cáo trực quan và xuất bản ra định dạng Excel (xlsx) hoặc Word (docx) chuyên nghiệp.")

# ----------------- KHI ĐÃ UPLOAD FILE -----------------
else:

    # Phân chia 2 cột giao diện: Cột trái (Văn bản gốc), Cột phải (Báo cáo Thẩm định)
    col_left, col_right = st.columns([1, 1.25], gap="large")

    # ================= CỘT TRÁI: HIỂN THỊ VĂN BẢN GỐC =================
    with col_left:
        st.subheader("📄 Văn bản gốc hợp đồng")
        
        # Ô tìm kiếm từ khóa
        search_query = st.text_input(
            "🔍 Tìm kiếm và tô sáng từ khóa:",
            placeholder="Nhập cụm từ cần tìm (ví dụ: bảo hành, phạt, lãi suất...)"
        )
        
        # Đọc dữ liệu từ session_state và hiển thị
        contract_html = ""
        if st.session_state.get("full_contract_text"):
            lines = st.session_state.full_contract_text.split('\n')
            for line in lines:
                text = line.strip()
                if not text:
                    continue
                    
                # Tô sáng nếu có từ khóa tìm kiếm
                if search_query:
                    escaped_query = re.escape(search_query)
                    pattern = re.compile(f"({escaped_query})", re.IGNORECASE)
                    text = pattern.sub(r'<span style="background-color: #fef08a; color: #854d0e; padding: 2px 4px; border-radius: 3px; font-weight: 600;">\1</span>', text)
                
                # Render các dòng đặc biệt của bảng
                if text.startswith("--- [BẮT ĐẦU BẢNG BIỂU] ---") or text.startswith("--- [KẾT THÚC BẢNG BIỂU] ---"):
                    contract_html += f"<p style='margin-bottom: 0.7rem; text-align: center; font-size: 13px; font-weight: bold; color: #64748b;'>{text}</p>"
                elif " | " in text:
                    contract_html += f"<p style='margin-bottom: 0.3rem; text-align: justify; font-size: 14px; background-color: #f8fafc; padding: 5px; border-bottom: 1px solid #e2e8f0;'>{text}</p>"
                else:
                    contract_html += f"<p style='margin-bottom: 0.7rem; text-align: justify; font-size: 14.5px;'>{text}</p>"
                
        # Render toàn bộ HTML văn bản gốc
        st.markdown(f'<div class="contract-viewer">{contract_html}</div>', unsafe_allow_html=True)
        st.caption(f"Tệp: {st.session_state.current_file_name} | Đang xem ở định dạng văn bản thô.")


    # ================= CỘT PHẢI: CHI TIẾT BÁO CÁO & BIÊN SOẠN BẰNG AI =================
    with col_right:
        st.subheader("🤖 Trợ lý Pháp lý AI (Gemini 3.5 Flash)")
        
        full_contract_text = st.session_state.get("full_contract_text", "")
        if not full_contract_text or full_contract_text.startswith("LỖI HỆ THỐNG"):
            st.error("Không thể trích xuất dữ liệu từ file này. Vui lòng kiểm tra lại định dạng file Word.")
        else:
            tab_appraisal, tab_edit = st.tabs(["📋 Báo cáo Thẩm định", "✍️ Chỉnh sửa & Biên soạn"])
            
            # Tab 1: Báo cáo Thẩm định
            with tab_appraisal:
                st.caption(f"Đã trích xuất thành công {len(full_contract_text)} ký tự văn bản.")
                
                if not api_key:
                    st.warning("⚠️ Vui lòng điền Gemini API Key ở Sidebar (thanh cấu hình bên trái) để kích hoạt tính năng này.")
                else:
                    col_run_ai, col_clear_ai = st.columns([2, 1])
                    with col_run_ai:
                        run_appraisal = st.button("🚀 Bắt đầu Thẩm định bằng AI", use_container_width=True, type="primary")
                    with col_clear_ai:
                        if st.session_state.gemini_report:
                            if st.button("🗑️ Xóa kết quả", use_container_width=True):
                                st.session_state.gemini_report = None
                                st.rerun()
                    
                    # Xử lý khi bấm nút thẩm định
                    if run_appraisal:
                        with st.spinner("Gemini 3.5 Flash đang đọc và phân tích toàn bộ hợp đồng (có thể mất 15-30 giây)..."):
                            selected_templates = st.session_state.get("selected_templates", [])
                            tieu_chuan_doi_chieu = load_templates(selected_templates)
                            ai_result = tham_dinh_bang_gemini(full_contract_text, api_key, tieu_chuan_doi_chieu)
                            st.session_state.gemini_report = ai_result
                            st.toast("Thẩm định bằng AI hoàn tất!")
                            st.rerun()
                    
                    # Hiển thị kết quả nếu đã có báo cáo
                    if st.session_state.gemini_report:
                        st.success("Báo cáo phân tích hợp đồng từ Gemini 3.5 Flash:")
                        
                        # Loại bỏ markdown tags thừa nếu có
                        report_content = st.session_state.gemini_report.strip()
                        if report_content.startswith("```markdown"):
                            report_content = report_content[11:].strip()
                        elif report_content.startswith("```"):
                            report_content = report_content[3:].strip()
                        if report_content.endswith("```"):
                            report_content = report_content[:-3].strip()
                        
                        # Hiển thị bảng Markdown lên giao diện
                        st.markdown(report_content)
                        
                        # Xuất báo cáo AI ra Word và PDF
                        st.divider()
                        st.write("Tải xuống báo cáo thẩm định:")
                        col_ai_word, col_ai_pdf = st.columns(2)
                        with col_ai_word:
                            ai_word_data = export_gemini_table_to_word(report_content, contract_name=st.session_state.current_file_name)
                            st.download_button(
                                label="📥 Tải Báo cáo dạng Word (.docx)",
                                data=ai_word_data,
                                file_name=f"Bao_cao_AI_Gemini_{datetime.now().strftime('%Y%m%d')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True,
                                key="btn_download_word_report"
                            )
                        with col_ai_pdf:
                            ai_pdf_data = export_gemini_table_to_pdf(report_content, contract_name=st.session_state.current_file_name)
                            st.download_button(
                                label="📥 Tải Báo cáo dạng PDF (.pdf)",
                                data=ai_pdf_data,
                                file_name=f"Bao_cao_AI_Gemini_{datetime.now().strftime('%Y%m%d')}.pdf",
                                mime="application/pdf",
                                use_container_width=True,
                                key="btn_download_pdf_report"
                            )
            
            # Tab 2: Chỉnh sửa & Biên soạn
            with tab_edit:
                st.caption(f"Yêu cầu AI viết lại hoặc tinh chỉnh nội dung các điều khoản trong hợp đồng.")
                
                if not api_key:
                    st.warning("⚠️ Vui lòng điền Gemini API Key ở Sidebar (thanh cấu hình bên trái) để kích hoạt tính năng này.")
                else:
                    # Gợi ý một số yêu cầu nhanh
                    quick_options = [
                        "Tự nhập yêu cầu chỉnh sửa...",
                        "Viết lại nội dung Điều 4 đảm bảo ngắn gọn súc tích, cô đọng.",
                        "Tối ưu hóa các điều khoản về quyền và nghĩa vụ của Bên Mua để chặt chẽ hơn.",
                        "Tăng mức phạt vi phạm hợp đồng và bổ sung bồi thường thiệt hại tối đa pháp luật cho phép.",
                        "Làm rõ điều khoản Bất khả kháng chi tiết, bổ sung các sự kiện như thiên tai, dịch bệnh.",
                        "Chỉnh sửa điều khoản Giải quyết tranh chấp theo hướng ưu tiên hòa giải trước khi ra Tòa án."
                    ]
                    
                    selected_quick = st.selectbox(
                        "💡 Chọn mẫu yêu cầu chỉnh sửa nhanh:",
                        options=quick_options,
                        index=0
                    )
                    
                    # Xác định giá trị mặc định cho text_area
                    default_text = ""
                    if selected_quick != "Tự nhập yêu cầu chỉnh sửa...":
                        default_text = selected_quick
                        
                    user_edit_req = st.text_area(
                        "✍️ Nhập chi tiết yêu cầu chỉnh sửa của bạn:",
                        value=default_text,
                        height=100,
                        placeholder="Ví dụ: Viết lại nội dung Điều 4 đảm bảo ngắn gọn súc tích, cô đọng..."
                    )
                    
                    col_run_edit, col_clear_edit = st.columns([2, 1])
                    with col_run_edit:
                        run_edit = st.button("🚀 Bắt đầu Chỉnh sửa bằng AI", use_container_width=True, type="primary", key="btn_run_edit")
                    with col_clear_edit:
                        if st.session_state.gemini_edit_result:
                            if st.button("🗑️ Xóa kết quả", use_container_width=True, key="btn_clear_edit"):
                                st.session_state.gemini_edit_result = None
                                st.rerun()
                                
                    if run_edit:
                        if not user_edit_req.strip():
                            st.warning("Vui lòng nhập nội dung yêu cầu chỉnh sửa trước khi chạy.")
                        else:
                            with st.spinner("Gemini 3.5 Flash đang biên tập và chỉnh sửa điều khoản theo yêu cầu..."):
                                edit_result = chinh_sua_hop_dong_bang_gemini(full_contract_text, api_key, user_edit_req)
                                st.session_state.gemini_edit_result = edit_result
                                st.toast("Chỉnh sửa bằng AI hoàn tất!")
                                st.rerun()
                                
                    if st.session_state.gemini_edit_result:
                        st.success("Đề xuất chỉnh sửa từ AI:")
                        
                        edit_content = st.session_state.gemini_edit_result.strip()
                        if edit_content.startswith("```markdown"):
                            edit_content = edit_content[11:].strip()
                        elif edit_content.startswith("```"):
                            edit_content = edit_content[3:].strip()
                        if edit_content.endswith("```"):
                            edit_content = edit_content[:-3].strip()
                            
                        st.markdown(edit_content)
                        
                        # Xuất báo cáo chỉnh sửa ra Word và PDF
                        st.divider()
                        st.write("Tải xuống đề xuất chỉnh sửa:")
                        col_edit_word, col_edit_pdf = st.columns(2)
                        with col_edit_word:
                            edit_word_data = export_gemini_table_to_word(edit_content, contract_name=f"De_xuat_chinh_sua_{st.session_state.current_file_name}")
                            st.download_button(
                                label="📥 Tải Đề xuất dạng Word (.docx)",
                                data=edit_word_data,
                                file_name=f"De_xuat_chinh_sua_{datetime.now().strftime('%Y%m%d')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True,
                                key="btn_download_word_edit"
                            )
                        with col_edit_pdf:
                            edit_pdf_data = export_gemini_table_to_pdf(edit_content, contract_name=f"De_xuat_chinh_sua_{st.session_state.current_file_name}")
                            st.download_button(
                                label="📥 Tải Đề xuất dạng PDF (.pdf)",
                                data=edit_pdf_data,
                                file_name=f"De_xuat_chinh_sua_{datetime.now().strftime('%Y%m%d')}.pdf",
                                mime="application/pdf",
                                use_container_width=True,
                                key="btn_download_pdf_edit"
                            )
