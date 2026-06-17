import io
import docx
from docx import Document
from docx.document import Document as DocumentClass
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

def iter_block_items(parent):
    """
    Duyệt qua các phần tử của tài liệu Word (đoạn văn và bảng biểu) theo đúng thứ tự xuất hiện.
    """
    if isinstance(parent, DocumentClass):
        parent_elm = parent.element.body
    else:
        raise ValueError("Đối tượng cha phải là một tài liệu Word (Document).")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def parse_docx(file_bytes):
    """
    Đọc file docx từ bytes và trả về danh sách các khối cấu trúc (paragraph hoặc table) theo thứ tự.
    """
    doc_file = io.BytesIO(file_bytes)
    doc = Document(doc_file)
    
    blocks = []
    
    for item in iter_block_items(doc):
        if isinstance(item, Paragraph):
            text = item.text.strip()
            # Bỏ qua dòng trống
            if not text:
                continue
                
            style_name = item.style.name if item.style else "Normal"
            is_heading = style_name.startswith("Heading") or style_name in ["Title", "Subtitle"]
            
            # Phân cấp tiêu đề dựa trên tên style hoặc tự động phát hiện
            heading_level = 0
            if is_heading:
                if style_name.startswith("Heading "):
                    try:
                        heading_level = int(style_name.replace("Heading ", ""))
                    except ValueError:
                        heading_level = 1
                elif style_name == "Title":
                    heading_level = 1
                elif style_name == "Subtitle":
                    heading_level = 2
            
            blocks.append({
                "type": "paragraph",
                "text": text,
                "style": style_name,
                "is_heading": is_heading,
                "heading_level": heading_level
            })
            
        elif isinstance(item, Table):
            table_data = []
            for row in item.rows:
                row_data = []
                for cell in row.cells:
                    # Lấy text trong ô, các đoạn văn phân tách bằng xuống dòng
                    cell_text = "\n".join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
                    row_data.append(cell_text)
                table_data.append(row_data)
            
            # Chỉ thêm bảng nếu có dữ liệu
            if table_data:
                blocks.append({
                    "type": "table",
                    "data": table_data
                })
                
    return blocks

def doc_hop_dong_toan_dien(file_input):
    """
    Đọc toàn bộ văn bản trong file .docx bao gồm văn bản thường và văn bản trong bảng.
    Chấp nhận cả đường dẫn file (str) hoặc đối tượng file upload từ Streamlit (BytesIO).
    """
    try:
        # Khởi tạo document
        doc = docx.Document(file_input)
        noi_dung_day_du = []
        
        # LƯU Ý BẮT bắt buộc: Duyệt qua từng phần tử theo thứ tự xuất hiện trong file
        # để tránh việc văn bản trong bảng bị đảo lộn thứ tự so với văn bản thường.
        for element in doc.element.body:
            # Nếu phần tử là một đoạn văn bản (Paragraph)
            if element.tag.endswith('p'):
                p = docx.text.paragraph.Paragraph(element, doc)
                if p.text.strip():
                    noi_dung_day_du.append(p.text.strip())
                    
            # Nếu phần tử là một bảng biểu (Table)
            elif element.tag.endswith('tbl'):
                t = docx.table.Table(element, doc)
                noi_dung_day_du.append("\n--- [BẮT ĐẦU BẢNG BIỂU] ---")
                
                for row in t.rows:
                    row_text = []
                    for cell in row.cells:
                        # Loại bỏ khoảng trắng thừa trong ô
                        cell_content = cell.text.strip()
                        # Tránh lặp từ đối với các ô bị hợp nhất (Merged cells)
                        if cell_content and cell_content not in row_text:
                            row_text.append(cell_content)
                    
                    if row_text:
                        # Ghép các cột trong một hàng bằng ký tự | để AI dễ hiểu cấu trúc bảng
                        noi_dung_day_du.append(" | ".join(row_text))
                        
                noi_dung_day_du.append("--- [KẾT THÚC BẢNG BIỂU] ---\n")
        
        # Ghép toàn bộ các dòng lại thành một chuỗi văn bản duy nhất
        van_ban_cuoi_cung = "\n".join(noi_dung_day_du)
        return van_ban_cuoi_cung.strip()
        
    except Exception as e:
        return f"LỖI HỆ THỐNG ĐỌC FILE: {str(e)}"
